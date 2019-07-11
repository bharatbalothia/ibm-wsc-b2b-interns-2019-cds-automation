from docx.api import Document
import re

# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
'''
filename  = "C:\\Users\\RajnishKumarVENDORRo\\PycharmProjects\\CDS\\Test files\\Questionaire\\IBM_AS2_Questionnaire_EuroGIODE(1).docx"
document = Document(filename)
table = document.tables[1]

# Data will be a list of rows represented as dictionaries
# containing each row's data.
data = []
for table in document.tables:
    #print(table.rows[0].cells[0].text)
    if  "Trading Partner’s Configuration" in table.rows[0].cells[0].text:
        info ={}
        for row in table.rows:
            cells = row.cells
            if 'Mandatory' in cells[0].text and cells[1].text == "":
                print(cells[0].text)

        print(info)

'''
def questcheck(filename):
    document = Document(filename)
    for table in document.tables:
        if "Trading Partner’s Configuration" in table.rows[0].cells[0].text or 'Outbound Communication' in table.rows[0].cells[0].text:
            info ={}
            as2_url = ""
            for row in table.rows:
                cells = row.cells
                if 'AS2 Server URL' in cells[0].text:
                    as2_url = cells[1].text
                if 'Mandatory' in cells[0].text and cells[1].text == "" and 'If SSL is used' not in cells[0].text:
                    yield cells[0].text.split("(")[0]
                elif 'If SSL is used' in cells[0].text and cells[1].text == "" and "https" in as2_url:
                    yield cells[0].text.split("(")[0]

'''
keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)

    # Establish the mapping based on the first row
    # headers; these will become the keys of our dictionary
    if i == 0:
        keys = tuple(text)
        continue

    # Construct a dictionary for this row, mapping
    # keys to values for this row
    row_data = dict(zip(keys, text))
    data.append(row_data)
'''