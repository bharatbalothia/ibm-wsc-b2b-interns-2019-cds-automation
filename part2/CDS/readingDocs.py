import docx
from docx.enum.style import WD_STYLE_TYPE
import difflib
import re
import Email
import os
import itertools as it


def getText(filename):
    doc = docx.Document(filename)
    for para in doc.paragraphs:
        if re.search('[a-zA-Z]', para.text):
            yield para.text

def getruns(filename):
    doc = docx.Document(filename)
    for para in doc.paragraphs:
        for runs in para.runs:
            print(runs.text)



def surveyAnalyser(survey_filename):
    survey_sent = "C:\\Users\\RajnishKumarVENDORRo\\PycharmProjects\\CDS\\Test files\Intro email Attachments\\CustomerABC_Corp_TP Survey.docx"
    survey_recived = "C:\\Users\\RajnishKumarVENDORRo\\PycharmProjects\\CDS\\Test files\Survey\\TP1_ABC_Corp_TP Survey.docx"

    fields = ['Company Name','Contact Name','Email address','Technical Contact Name','Email address','Production EDI','Qualif','Test EDI']

    doc = docx.Document(survey_filename)
    fullText = []

    styles = doc.styles
    paragraph_styles = [
        s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
    ]

    '''
    for style in paragraph_styles:
        print(style.name)
        if style.name == 'List Paragraph':
            print("I'm a bullet")
    
    print(doc.styles)
    '''
    '''
    info = {}
    #print(getText(survey_recived))
    #print(getText(survey_recived))
    for lines in getText(survey_recived):
        if any(s in lines for s in fields):
            resp = lines.split(":")
            if len(resp) >1:
                info[resp[0]] = re.sub('_', '', resp[1])
    
    
    print(info)
    '''

    quest_map = {
        "AS2" : "IBM_AS2_Questionnaire_EuroGIODE(1).docx",
        "FTP(S)" : "IBM_FTP_Questionnaire_EuroGIODE.docx",
        "OFTP2" : "IBM_OFTP2_Questionnaire_EuroGIODE.docx"
    }
    comm = []
    flag =0
    for lines in getText(survey_filename):
        if "method of communication" in lines:
            flag = 1
            continue
        if flag == 1 and lines.strip().startswith("X"):
            print(lines.split())
            comm.append(lines.split()[1])

    files =[]
    for data in comm:
        name = quest_map[data]
        file_path = "C:\\Users\\RajnishKumarVENDORRo\\PycharmProjects\\CDS\Test files\\Questionaire\\" +name
        files.append(file_path)

    attachment = it.takewhile(lambda x: os.path.exists(x), files)
    return attachment
